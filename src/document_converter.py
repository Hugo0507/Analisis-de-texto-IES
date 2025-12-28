import os
import pandas as pd
from docx import Document
import PyPDF2
import pdfplumber
from pdfminer.high_level import extract_text
from collections import defaultdict
from src.utils.logger import get_logger

# Inicializar logger
logger = get_logger(__name__)


class DocumentConverter:
   

    def __init__(self):
        
        self.supported_formats = {
            '.txt': self.convert_txt,
            '.pdf': self.convert_pdf,
            '.docx': self.convert_docx,
            '.doc': self.convert_doc,
        }

        self.conversion_stats = {
            'total': 0,
            'successful': 0,
            'failed': 0,
            'by_format': defaultdict(lambda: {'success': 0, 'failed': 0}),
            'errors': []
        }

    def convert_txt(self, file_path, encoding='utf-8'):
  
        try:
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Error leyendo TXT: {e}")

    def convert_pdf(self, file_path):

        text = ""

        # 1ra opción: pdfminer.six (más robusto y preciso)
        try:
            text = extract_text(file_path)
            if text and text.strip():
                logger.debug(f"PDF extraído exitosamente con pdfminer.six: {file_path}")
                return text
        except Exception as e:
            logger.warning(f"pdfminer.six falló para {file_path}: {e}, intentando con PyPDF2...")

        # 2da opción: PyPDF2
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            if text.strip():
                logger.debug(f"PDF extraído exitosamente con PyPDF2: {file_path}")
                return text

        except Exception as e:
            logger.warning(f"PyPDF2 falló para {file_path}: {e}, intentando con pdfplumber...")

        # 3ra opción: pdfplumber (último recurso)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            if text.strip():
                return text
            else:
                raise Exception("No se pudo extraer texto del PDF")

        except Exception as e:
            raise Exception(f"Error leyendo PDF con todos los métodos: {e}")

    def convert_pdf_from_bytes(self, file_bytes, file_name="archivo.pdf"):

        text = ""

        # 1ra opción: pdfminer.six (más robusto y preciso)
        try:
            file_bytes.seek(0)
            text = extract_text(file_bytes)
            if text and text.strip():
                logger.debug(f"PDF extraído exitosamente con pdfminer.six desde memoria: {file_name}")
                return text
        except Exception as e:
            logger.warning(f"pdfminer.six falló para {file_name}: {e}, intentando con PyPDF2...")

        # 2da opción: PyPDF2
        try:
            file_bytes.seek(0)
            pdf_reader = PyPDF2.PdfReader(file_bytes)

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            if text.strip():
                logger.debug(f"PDF extraído exitosamente con PyPDF2 desde memoria: {file_name}")
                return text

        except Exception as e:
            logger.warning(f"PyPDF2 falló para {file_name}: {e}, intentando con pdfplumber...")

        # 3ra opción: pdfplumber (último recurso)
        try:
            file_bytes.seek(0)
            with pdfplumber.open(file_bytes) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            if text.strip():
                return text
            else:
                raise Exception("No se pudo extraer texto del PDF")

        except Exception as e:
            raise Exception(f"Error leyendo PDF desde memoria con todos los métodos: {e}")

    def convert_docx(self, file_path):

        try:
            doc = Document(file_path)
            text = ""

            # Extraer texto de párrafos
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"

            if text.strip():
                return text
            else:
                raise Exception("El documento está vacío")

        except Exception as e:
            raise Exception(f"Error leyendo DOCX: {e}")

    def convert_doc(self, file_path):

        raise Exception("Formato .doc no soportado directamente. Convierte a .docx primero.")

    def convert_from_bytes(self, file_bytes, file_name, file_extension):

        self.conversion_stats['total'] += 1

        result = {
            'file': file_name,
            'extension': file_extension,
            'success': False,
            'output': None,
            'text_length': 0,
            'error': None,
            'text': None
        }

        try:
            # Verificar formato soportado
            if file_extension == '.pdf':
                text = self.convert_pdf_from_bytes(file_bytes, file_name)
            elif file_extension == '.txt':
                file_bytes.seek(0)
                text = file_bytes.read().decode('utf-8', errors='ignore')
            elif file_extension == '.docx':
                # Para DOCX desde memoria, necesitamos python-docx que acepta BytesIO
                from docx import Document
                file_bytes.seek(0)
                doc = Document(file_bytes)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
            else:
                raise Exception(f"Formato {file_extension} no soportado")

            if not text or len(text.strip()) == 0:
                raise Exception("No se extrajo texto del documento")

            result['text'] = text
            result['success'] = True
            result['text_length'] = len(text)

            # Actualizar estadísticas
            self.conversion_stats['successful'] += 1
            self.conversion_stats['by_format'][file_extension]['success'] += 1

        except Exception as e:
            result['error'] = str(e)
            self.conversion_stats['failed'] += 1
            self.conversion_stats['by_format'][file_extension]['failed'] += 1
            self.conversion_stats['errors'].append({
                'file': file_name,
                'error': str(e)
            })

        return result

    def convert_file(self, file_path, output_path=None, encoding='utf-8'):

        self.conversion_stats['total'] += 1

        # Obtener extensión
        file_ext = os.path.splitext(file_path)[1].lower()

        result = {
            'file': file_path,
            'extension': file_ext,
            'success': False,
            'output': None,
            'text_length': 0,
            'error': None
        }

        try:
            # Verificar si el formato es soportado
            if file_ext not in self.supported_formats:
                raise Exception(f"Formato {file_ext} no soportado")

            # Convertir
            converter_func = self.supported_formats[file_ext]
            if file_ext == '.txt':
                text = converter_func(file_path, encoding)
            else:
                text = converter_func(file_path)

            if not text or len(text.strip()) == 0:
                raise Exception("No se extrajo texto del documento")

            # Guardar si se especificó ruta de salida
            if output_path:
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                with open(output_path, 'w', encoding=encoding, errors='ignore') as f:
                    f.write(text)
                result['output'] = output_path
            else:
                result['text'] = text

            result['success'] = True
            result['text_length'] = len(text)

            # Actualizar estadísticas
            self.conversion_stats['successful'] += 1
            self.conversion_stats['by_format'][file_ext]['success'] += 1

        except Exception as e:
            result['error'] = str(e)
            self.conversion_stats['failed'] += 1
            self.conversion_stats['by_format'][file_ext]['failed'] += 1
            self.conversion_stats['errors'].append({
                'file': file_path,
                'error': str(e)
            })

        return result

    def convert_batch(self, file_paths, output_folder=None, encoding='utf-8',
                      preserve_structure=True):

        results = []

        for file_path in file_paths:
            # Determinar ruta de salida
            output_path = None
            if output_folder:
                if preserve_structure:
                    # Mantener estructura de carpetas
                    rel_path = os.path.relpath(file_path)
                    base_name = os.path.splitext(rel_path)[0]
                    output_path = os.path.join(output_folder, base_name + '.txt')
                else:
                    # Todo en la misma carpeta
                    file_name = os.path.splitext(os.path.basename(file_path))[0]
                    output_path = os.path.join(output_folder, file_name + '.txt')

            result = self.convert_file(file_path, output_path, encoding)
            results.append(result)

        return results

    def get_conversion_statistics(self):
 
        stats = dict(self.conversion_stats)

        # Calcular tasas de éxito
        if stats['total'] > 0:
            stats['success_rate'] = round((stats['successful'] / stats['total']) * 100, 2)
            stats['failure_rate'] = round((stats['failed'] / stats['total']) * 100, 2)
        else:
            stats['success_rate'] = 0
            stats['failure_rate'] = 0

        return stats

    def create_conversion_report(self, results):

        data = []
        for result in results:
            data.append({
                'Archivo': os.path.basename(result['file']),
                'Ruta': result['file'],
                'Formato': result['extension'],
                'Estado': '✓ Exitoso' if result['success'] else '✗ Fallido',
                'Longitud Texto': result['text_length'],
                'Archivo Salida': result.get('output', 'N/A'),
                'Error': result.get('error', '-')
            })

        return pd.DataFrame(data)

    def create_format_summary(self, results):

        format_stats = defaultdict(lambda: {'total': 0, 'exitosos': 0, 'fallidos': 0})

        for result in results:
            fmt = result['extension']
            format_stats[fmt]['total'] += 1
            if result['success']:
                format_stats[fmt]['exitosos'] += 1
            else:
                format_stats[fmt]['fallidos'] += 1

        data = []
        for fmt, stats in format_stats.items():
            success_rate = (stats['exitosos'] / stats['total'] * 100) if stats['total'] > 0 else 0
            data.append({
                'Formato': fmt,
                'Total': stats['total'],
                'Exitosos': stats['exitosos'],
                'Fallidos': stats['fallidos'],
                'Tasa de Éxito (%)': round(success_rate, 2)
            })

        df = pd.DataFrame(data)
        df = df.sort_values('Total', ascending=False)

        return df

    def get_failed_conversions(self, results):

        failed = [r for r in results if not r['success']]

        data = []
        for result in failed:
            data.append({
                'Archivo': os.path.basename(result['file']),
                'Ruta': result['file'],
                'Formato': result['extension'],
                'Error': result.get('error', 'Desconocido')
            })

        return pd.DataFrame(data)

    def reset_statistics(self):
        
        self.conversion_stats = {
            'total': 0,
            'successful': 0,
            'failed': 0,
            'by_format': defaultdict(lambda: {'success': 0, 'failed': 0}),
            'errors': []
        }

    def validate_text_extraction(self, file_path, min_length=100):

        try:
            result = self.convert_file(file_path)

            if not result['success']:
                return {
                    'valid': False,
                    'reason': result.get('error', 'Error desconocido'),
                    'length': 0
                }

            text_length = result['text_length']

            if text_length < min_length:
                return {
                    'valid': False,
                    'reason': f'Texto muy corto ({text_length} caracteres, mínimo {min_length})',
                    'length': text_length
                }

            return {
                'valid': True,
                'length': text_length,
                'message': 'Validación exitosa'
            }

        except Exception as e:
            return {
                'valid': False,
                'reason': str(e),
                'length': 0
            }
